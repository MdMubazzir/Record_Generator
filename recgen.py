from flask import Flask, render_template, request, send_file, jsonify
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, nsmap
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
import webbrowser
from threading import Timer
import json
import base64
import io
import uuid
import os



nsmap['v'] = 'urn:schemas-microsoft-com:vml'

app = Flask(__name__, template_folder="templates")

app.config['MAX_CONTENT_LENGTH'] = 200 * 1024 * 1024

# Storage for templates (in production, use a database)
TEMPLATE_STORAGE = {}
STORAGE_FILE = 'templates.json'


# Load existing templates on startup
def add_footer_to_section(section, item):
    """Add footer text to a specific section (last page only)"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()

    # Always left align
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Main footer text
    run = footer_para.add_run(str(item.get("text", "")))
    run.font.name = item.get("font", "Calibri")
    run.font.size = Pt(int(item.get("size", 10)))
    run.font.bold = item.get("bold", False)

    # Force black color using XML
    rPr = run._r.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')
    rPr.append(color)

    # Add text block if enabled
    if item.get("text_enabled") and item.get("text_content"):
        text_para = footer.add_paragraph()
        text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        text_run = text_para.add_run(str(item.get("text_content", "")))
        text_run.font.name = item.get("text_font", "Calibri")
        text_run.font.size = Pt(int(item.get("text_size", 12)))

        # Force black color
        rPr2 = text_run._r.get_or_add_rPr()
        color2 = OxmlElement('w:color')
        color2.set(qn('w:val'), '000000')
        rPr2.append(color2)

def load_templates():
    global TEMPLATE_STORAGE
    if os.path.exists(STORAGE_FILE):
        try:
            with open(STORAGE_FILE, 'r') as f:
                TEMPLATE_STORAGE = json.load(f)
        except:
            TEMPLATE_STORAGE = {}


# Save templates to file
def save_templates():
    with open(STORAGE_FILE, 'w') as f:
        json.dump(TEMPLATE_STORAGE, f)


load_templates()


# ---------------- BORDER ----------------
def add_page_border(section):
    xml = parse_xml(
        r'<w:pgBorders %s w:offsetFrom="page">'
        r'<w:top w:val="single" w:sz="12" w:space="24"/>'
        r'<w:left w:val="single" w:sz="12" w:space="24"/>'
        r'<w:bottom w:val="single" w:sz="12" w:space="24"/>'
        r'<w:right w:val="single" w:sz="12" w:space="24"/>'
        r'</w:pgBorders>' % nsdecls('w')
    )
    section._sectPr.append(xml)


# ---------------- WATERMARK ----------------
def add_watermark(section, text):
    from docx.oxml import OxmlElement

    header = section.header
    paragraph = header.paragraphs[0]
    run = paragraph.add_run()

    pict = OxmlElement('w:pict')

    shape = OxmlElement('v:shape')
    shape.set('id', 'Watermark')
    shape.set('type', '#_x0000_t136')
    shape.set(
        'style',
        'position:absolute;'
        'width:500pt;height:120pt;'
        'rotation:315;'
        'mso-position-horizontal:center;'
        'mso-position-horizontal-relative:page;'
        'mso-position-vertical:center;'
        'mso-position-vertical-relative:page;'
    )
    shape.set('stroked', 'f')
    shape.set('fillcolor', '#d9d9d9')

    textpath = OxmlElement('v:textpath')
    textpath.set('style', "font-family:'Calibri';font-size:72pt")
    textpath.set('string', text)

    shape.append(textpath)
    pict.append(shape)
    run._r.append(pict)


# ---------------- MARGINS ----------------
def set_narrow_margins(section):
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)


# ---------------- FOOTER ----------------
def add_footer(section, text, font="Calibri", size=10, alignment="left",
               before_lines=0, spacing=0, text_enabled=False, text_content="",
               text_font="Calibri", text_size=12, bold=False):
    """Add footer text to the section with advanced options"""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.clear()

    # Always left align
    footer_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Main footer text
    run = footer_para.add_run(str(text))
    run.font.name = font
    run.font.size = Pt(int(size))
    run.font.bold = bold

    # Force black color using XML
    rPr = run._r.get_or_add_rPr()
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '000000')  # Black hex code
    rPr.append(color)

    # Add text block if enabled
    if text_enabled and text_content:
        text_para = footer.add_paragraph()
        text_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        text_run = text_para.add_run(str(text_content))
        text_run.font.name = text_font
        text_run.font.size = Pt(int(text_size))

        # Force black color using XML
        rPr2 = text_run._r.get_or_add_rPr()
        color2 = OxmlElement('w:color')
        color2.set(qn('w:val'), '000000')  # Black hex code
        rPr2.append(color2)

    # Add after lines
    for _ in range(int(spacing)):
        footer.add_paragraph()



# ---------------- HEADINGS + TEXT + IMAGE ENGINE ----------------

def add_headings(doc, headings_data):
    """Process only regular headings (no footers)"""
    print(f"Processing {len(headings_data)} headings")

    for idx, item in enumerate(headings_data):
        print(f"\n--- Heading {idx + 1} ---")

        # ---------- BEFORE LINES ----------
        try:
            before_lines = int(item.get("before_lines", 0))
            for _ in range(before_lines):
                doc.add_paragraph("")
        except:
            pass

        # ---------- HEADING ----------
        p = doc.add_paragraph()
        run = p.add_run(str(item.get("text", "")))
        run.font.name = item.get("font", "Calibri")
        run.font.size = Pt(int(item.get("size", 16)))
        run.font.bold = item.get("bold", False)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # Handle spacing after heading
        spacing = item.get("spacing", "1")
        if spacing == "till_end":
            doc.add_page_break()
        else:
            try:
                for _ in range(int(spacing)):
                    doc.add_paragraph("")
            except:
                pass
        # ---------- TEXT BLOCK ----------
        if item.get("text_enabled"):
            print("Adding text block...")

            try:
                before = int(item.get("text_before", 0))
                for _ in range(before):
                    doc.add_paragraph("")
            except:
                pass

            text_content = item.get("text_content", "")
            if text_content:
                tp = doc.add_paragraph()
                tr = tp.add_run(str(text_content))
                tr.font.name = item.get("text_font", "Calibri")
                tr.font.size = Pt(int(item.get("text_size", 12)))

            if item.get("text_till_end"):
                doc.add_page_break()
            else:
                try:
                    after = int(item.get("text_after", 0))
                    for _ in range(after):
                        doc.add_paragraph("")
                except:
                    pass


        # ---------- IMAGE BLOCK ----------
        if item.get("image_enabled") and item.get("image_data"):
            print("Adding image block...")
            try:
                image_data_str = item.get("image_data", "")

                if image_data_str:
                    if ',' in image_data_str:
                        image_data_str = image_data_str.split(',')[1]

                    image_data = base64.b64decode(image_data_str)
                    image_stream = io.BytesIO(image_data)

                    width_inches = float(item.get("image_width", 4))
                    height_inches = float(item.get("image_height", 3))

                    pic = doc.add_picture(image_stream, width=Inches(width_inches))
                    pic.height = Inches(height_inches)

                    try:
                        img_after = int(item.get("image_after", 0))
                        for _ in range(img_after):
                            doc.add_paragraph("")
                    except:
                        pass

                    if item.get("image_till_end"):
                        doc.add_page_break()

                    print("Image added successfully")
            except Exception as e:
                print(f"Error adding image: {str(e)}")
# ---------------- SAVE TEMPLATE ----------------
@app.route("/save_template", methods=["POST"])
def save_template():
    try:
        data = request.json
        template_code = str(uuid.uuid4())[:8].upper()

        # Store template with metadata
        TEMPLATE_STORAGE[template_code] = {
            'headings': data.get('headings', []),
            'border': data.get('border', False),
            'watermark': data.get('watermark', ''),
            'layout': data.get('layout', 'narrow'),
            'created_at': str(uuid.uuid1()),
            'heading_count': data.get('heading_count', 0)
        }

        save_templates()

        return jsonify({
            'success': True,
            'code': template_code
        })
    except Exception as e:
        return jsonify({
            'success': False,
            'error': str(e)
        })


# ---------------- LOAD TEMPLATE ----------------
@app.route("/load_template/<code>", methods=["GET"])
def load_template(code):
    code = code.upper().strip()

    if code in TEMPLATE_STORAGE:
        return jsonify({
            'success': True,
            'template': TEMPLATE_STORAGE[code]
        })
    else:
        return jsonify({
            'success': False,
            'error': 'Template not found'
        })


# ---------------- MAIN ROUTE ----------------
# ---------------- MAIN ROUTE ----------------
@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        print("\n===== FORM SUBMITTED =====")

        border_enabled = request.form.get("border")
        watermark_text = request.form.get("watermark")

        headings_json = request.form.get("headings", "[]")

        try:
            headings_data = json.loads(headings_json)
        except Exception as e:
            print(f"Error parsing JSON: {e}")
            headings_data = []

        # Separate regular headings from footers
        regular_headings = []
        footer_items = []

        for item in headings_data:
            if item.get("is_footer"):
                footer_items.append(item)
            else:
                regular_headings.append(item)

        doc = Document()
        section = doc.sections[0]
        layout = request.form.get("layout", "narrow")
        if layout == "narrow":
            set_narrow_margins(section)
        else:
            # Normal margins (1 inch)
            section.top_margin = Inches(1.0)
            section.bottom_margin = Inches(1.0)
            section.left_margin = Inches(1.0)
            section.right_margin = Inches(1.0)

        if border_enabled:
            add_page_border(section)

        if watermark_text:
            add_watermark(section, watermark_text)

        # If we have footers, we need to split content so last heading shares page with footer
        if footer_items and regular_headings:
            # Calculate total lines to find which page the last heading will be on
            LINES_PER_PAGE = 45  # Approximate lines per page

            total_lines = 0
            for i, item in enumerate(regular_headings):
                # Lines before this heading
                total_lines += int(item.get("before_lines", 0))
                # The heading itself (1 line)
                total_lines += 1

                # Spacing after heading
                if item.get("spacing") == "till_end":
                    # Push to next page
                    total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                else:
                    total_lines += int(item.get("spacing", 1))

                # Text block lines
                if item.get("text_enabled"):
                    total_lines += int(item.get("text_before", 0))
                    text_content = item.get("text_content", "")
                    # Estimate wrapped lines (80 chars per line)
                    if text_content:
                        wrapped_lines = sum(max(1, (len(line) // 80) + 1) for line in text_content.split('\n'))
                        total_lines += wrapped_lines

                    if item.get("text_till_end"):
                        total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                    else:
                        total_lines += int(item.get("text_after", 0))

                # Image lines (approx 6 lines per inch of height)
                if item.get("image_enabled") and item.get("image_data"):
                    img_height = float(item.get("image_height", 3))
                    total_lines += int(img_height * 6)

                    if item.get("image_till_end"):
                        total_lines = ((total_lines // LINES_PER_PAGE) + 1) * LINES_PER_PAGE
                    else:
                        total_lines += int(item.get("image_after", 0))

            # Find which page the last heading will be on
            last_heading_page = (total_lines // LINES_PER_PAGE) + 1
            print(f"Last heading will be on page: {last_heading_page}")

            # IF EVERYTHING FITS ON ONE PAGE: No section break needed
            if last_heading_page == 1:
                add_headings(doc, regular_headings)
                final_section = doc.sections[-1]
                final_section.footer.is_linked_to_previous = False
                if final_section.footer.paragraphs:
                    final_section.footer.paragraphs[0].clear()
                for footer_item in footer_items:
                    add_footer_to_section(final_section, footer_item)

            # If we have multiple headings and content spans multiple pages
            elif len(regular_headings) > 1:
                main_headings = regular_headings[:-1]
                final_heading = [regular_headings[-1]]

                # Process main headings
                add_headings(doc, main_headings)

                # Add section break for the LAST page (where footer will be)
                doc.add_section(WD_SECTION.NEW_PAGE)
                final_section = doc.sections[-1]

                # Copy settings
                set_narrow_margins(final_section)
                if border_enabled:
                    add_page_border(final_section)
                if watermark_text:
                    add_watermark(final_section, watermark_text)

                # Unlink footer so we can add it to this section only
                final_section.footer.is_linked_to_previous = False
                if final_section.footer.paragraphs:
                    final_section.footer.paragraphs[0].clear()

                # Add the LAST heading to this section
                add_headings(doc, final_heading)

                # Add footer to SAME section (same page as last heading)
                for footer_item in footer_items:
                    add_footer_to_section(final_section, footer_item)

            else:
                # Only one heading - it shares the page with footer
                # Just unlink footer in current section and add both
                section.footer.is_linked_to_previous = False
                if section.footer.paragraphs:
                    section.footer.paragraphs[0].clear()

                # Add the single heading
                add_headings(doc, regular_headings)

                # Add footer to same section (same page)
                for footer_item in footer_items:
                    add_footer_to_section(section, footer_item)

        elif footer_items and not regular_headings:
            # Only footer, no headings - footer is the only content
            section.footer.is_linked_to_previous = False
            if section.footer.paragraphs:
                section.footer.paragraphs[0].clear()

            for footer_item in footer_items:
                add_footer_to_section(section, footer_item)

        else:
            # No footers - normal processing
            add_headings(doc, regular_headings)
        filename = "record.docx"
        doc.save(filename)

        return send_file(filename, as_attachment=True)

    return render_template("index.html")

# ---------------- RUN ----------------


if __name__ == "__main__":
    app.run(debug=False)
